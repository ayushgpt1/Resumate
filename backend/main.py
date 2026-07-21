import os
from dotenv import load_dotenv
# Load environment variables from this backend directory first to prevent import-order issues.
# override=True ensures backend/.env wins over any stale shell/system env vars.
load_dotenv(os.path.join(os.path.dirname(__file__), ".env"), override=True)

import json
import pdfplumber
import docx
import pytesseract
import re
import pandas as pd
import traceback
import time
import random
from PIL import Image
import flask
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from groq import Groq
import google.generativeai as genai
from controller.auth import auth_required
from models.candidate import CandidateProcessor
from utils.supabase_client import initialize_supabase
from utils.supabase_storage import upload_resume, ensure_bucket_exists
from utils.supabase_db import store_candidate, store_job_role, get_candidates_for_role, get_all_roles, delete_candidate, remove_candidate_from_role
from utils.local_storage import delete_candidate_local, delete_role_local
from linkedin import register_linkedin_routes

def get_tesseract_path():
    possible_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'/usr/bin/tesseract',
        r'/usr/local/bin/tesseract',
    ]
    env_path = os.getenv("TESSERACT_PATH")
    if env_path:
        possible_paths.insert(0, env_path)
    for path in possible_paths:
        if os.path.exists(path):
            return path
    return r'C:\Program Files\Tesseract-OCR\tesseract.exe'

pytesseract.pytesseract.tesseract_cmd = get_tesseract_path()

GROQ_API_KEY = os.getenv("GROQ_API_KEY")
QWEN_MODEL_NAME = os.getenv("QWEN_MODEL_NAME", "llama-3.3-70b-versatile")
groq_client = None
if GROQ_API_KEY:
    try:
        groq_client = Groq(api_key=GROQ_API_KEY)
        print("✓ Groq API key found. Using Groq as primary AI provider.")
    except Exception as e:
        print(f"✗ ERROR: Failed to initialize Groq client: {e}")
        groq_client = None
else:
    print("✗ GROQ_API_KEY not found. Groq will not be available.")

# Initialize Gemini for email templating
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
if GEMINI_API_KEY.startswith('"') and GEMINI_API_KEY.endswith('"'):
    GEMINI_API_KEY = GEMINI_API_KEY[1:-1]
if GEMINI_API_KEY.startswith("'") and GEMINI_API_KEY.endswith("'"):
    GEMINI_API_KEY = GEMINI_API_KEY[1:-1]
genai_client = None
if GEMINI_API_KEY:
    try:
        genai.configure(api_key=GEMINI_API_KEY)
        genai_client = True
        print("✓ Gemini API key found. Using Gemini for email templating.")
    except Exception as e:
        print(f"✗ Gemini initialization failed: {e}")
        genai_client = None
else:
    print("✗ GEMINI_API_KEY not found. Gemini will not be available for email templating.")

app = Flask(__name__)
CORS(app, supports_credentials=True)
supabase = initialize_supabase()

if supabase:
    from utils.supabase_db import test_supabase_connection
    db_ok, db_msg = test_supabase_connection()
    if db_ok:
        print(f"✓ Supabase database connection: {db_msg}")
    else:
        print(f"✗ Supabase database connection failed: {db_msg}")
else:
    print("✗ Supabase client not initialized.")

try:
    ensure_bucket_exists()
except Exception as e:
    print(f"Warning: Could not initialize storage bucket at startup: {e}")

app.config['UPLOAD_FOLDER'] = 'uploads'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

register_linkedin_routes(app)

@app.route('/uploads/<path:filename>', methods=['GET'])
def serve_upload(filename):
    safe_path = os.path.normpath(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    if not safe_path.startswith(os.path.normpath(app.config['UPLOAD_FOLDER'])):
        return jsonify({"error": "Invalid path"}), 403
    if not os.path.exists(safe_path):
        return jsonify({"error": "File not found"}), 404
    return flask.send_file(safe_path, as_attachment=False)

@app.route('/', methods=['GET'])
def home():
    return jsonify({
        "api": "HR Candidate Search Platform",
        "version": "1.0.0",
        "endpoints": {
            "resume_processing": ["/process", "/roles", "/candidates_for_role/{role_name}"],
            "candidate_search": ["/linkedin/candidates/search", "/linkedin/candidates/{id}", "/linkedin/requirements/submit"]
        }
    })

def extract_text_from_pdf(filepath):
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    except:
        pass
    return text.strip()

def extract_text_from_docx(filepath):
    try:
        doc = docx.Document(filepath)
        return "\n".join([para.text for para in doc.paragraphs]).strip()
    except:
        return ""

def extract_text_from_doc(filepath):
    try:
        import olefile
        ole = olefile.OleFileIO(filepath)
        if ole.exists('WordDocument'):
            doc_stream = ole.openstream('WordDocument')
            data = doc_stream.read()
            ole.close()
            best_text = ""
            best_score = 0
            for encoding in ['utf-16le', 'utf-8', 'latin-1']:
                try:
                    raw = data.decode(encoding, errors='replace')
                    cleaned = ''.join(c for c in raw if c.isprintable() or c in '\n\r\t ')
                    if cleaned.strip():
                        alpha = sum(1 for c in cleaned if c.isascii() and c.isalnum())
                        total = len(cleaned.strip())
                        score = alpha / total if total > 0 else 0
                        if score > best_score:
                            best_score = score
                            best_text = cleaned
                except:
                    pass
            if best_text and best_score > 0.3 and len(best_text.strip()) > 50:
                lines = best_text.split('\n')
                clean_lines = []
                for line in lines:
                    stripped = line.strip()
                    if stripped:
                        alpha_ratio = sum(1 for c in stripped if c.isascii() and c.isalnum()) / len(stripped)
                        if alpha_ratio > 0.4:
                            clean_lines.append(stripped)
                result = '\n'.join(clean_lines)
                if len(result) > 50:
                    return result
        else:
            ole.close()
    except ImportError:
        print("  - olefile not installed. Install with: pip install olefile")
    except:
        pass
    try:
        doc = docx.Document(filepath)
        text = "\n".join([para.text for para in doc.paragraphs]).strip()
        if text:
            return text
    except:
        pass
    return ""

def extract_text_from_image(filepath):
    try:
        img = Image.open(filepath)
        text = pytesseract.image_to_string(img).strip()
        if text:
            print(f"  - Extracted {len(text)} chars from image using Tesseract")
            return text
        else:
            print(f"  - Tesseract OCR returned empty text for image")
            return ""
    except Exception as tesseract_error:
        print(f"  - Tesseract OCR failed ({tesseract_error})")
        return ""

def extract_text_from_excel(filepath):
    try:
        text_parts = []
        try:
            excel_file = pd.ExcelFile(filepath)
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(filepath, sheet_name=sheet_name)
                text_parts.append(f"Sheet: {sheet_name}\n{df.to_string(index=False)}")
            result = "\n\n".join(text_parts).strip()
            if result:
                return result
        except:
            pass
        try:
            import openpyxl
            workbook = openpyxl.load_workbook(filepath, data_only=True)
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = []
                for row in sheet.iter_rows():
                    row_values = []
                    for cell in row:
                        value = cell.value
                        if value is not None:
                            row_values.append(str(value))
                    if row_values:
                        sheet_text.append(" | ".join(row_values))
                if sheet_text:
                    text_parts.append(f"Sheet: {sheet_name}\n" + "\n".join(sheet_text))
            return "\n\n".join(text_parts).strip()
        except:
            pass
        try:
            with open(filepath, 'rb') as f:
                binary_content = f.read()
                text_content = re.findall(b'[a-zA-Z0-9 .,:;\'"-]{4,}', binary_content)
                if text_content:
                    decoded_text = [t.decode('utf-8', errors='ignore') for t in text_content]
                    return "\n".join(decoded_text)
        except:
            pass
        return "Unable to extract text from Excel file."
    except Exception as e:
        return f"Failed to extract text from Excel file: {str(e)}"

def extract_resume_text(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext == ".docx":
        return extract_text_from_docx(filepath)
    elif ext == ".doc":
        return extract_text_from_doc(filepath)
    elif ext in [".png", ".jpg", ".jpeg"]:
        return extract_text_from_image(filepath)
    elif ext in [".xls", ".xlsx"]:
        return extract_text_from_excel(filepath)
    else:
        return ""

def fix_incomplete_json(json_str):
    try:
        json.loads(json_str)
        return json_str
    except json.JSONDecodeError:
        open_braces = json_str.count('{')
        close_braces = json_str.count('}')
        if open_braces > close_braces:
            json_str += '}' * (open_braces - close_braces)
        if '"explanation":' not in json_str and '"ranking_score":' in json_str:
            json_str = json_str.rstrip('}')
            if json_str.endswith(','):
                json_str += ' "explanation": "No explanation provided."}'
            else:
                json_str += ', "explanation": "No explanation provided."}'
        try:
            json.loads(json_str)
            return json_str
        except:
            match = re.search(r'\{.*\}', json_str, re.DOTALL)
            if match:
                potential_json = match.group(0)
                try:
                    json.loads(potential_json)
                    return potential_json
                except:
                    pass
            return """{"name": "Error Processing", "email": "Not provided", "phone_number": "Not provided", "experience": "Error occurred during processing", "skills": [], "cultural_fit": "Unknown", "communication": "Unknown", "ranking_score": 0, "explanation": "An error occurred during AI processing of this resume."}"""

def call_groq(prompt):
    try:
        print(f"Calling Groq API with model: {QWEN_MODEL_NAME}")
        if len(prompt) > 32000:
            print(f"Warning: Prompt length ({len(prompt)} chars) may exceed token limits.")
        completion = groq_client.chat.completions.create(
            model=QWEN_MODEL_NAME,
            messages=[
                {
                    "role": "system",
                    "content": """You are an HR evaluator. Return ONLY valid JSON output with:
1. No markdown code blocks (no ```json or ```)
2. No additional text outside the JSON object
3. No thinking process or commentary
4. Properly formatted JSON with all strings quoted
5. Make sure all fields including 'explanation' are included and all brackets are properly closed"""
                },
                {"role": "user", "content": prompt}
            ],
            response_format={"type": "json_object"}
        )
        response = completion.choices[0].message.content
        response = re.sub(r'```json|```', '', response).strip()
        try:
            json.loads(response)
            return response
        except json.JSONDecodeError:
            match = re.search(r'\{.*\}', response, re.DOTALL)
            if match:
                extracted = match.group(0)
                try:
                    json.loads(extracted)
                    print(f"Extracted valid JSON from Groq response")
                    return extracted
                except:
                    pass
            fixed_response = fix_incomplete_json(response)
            print(f"Fixed incomplete JSON from Groq: {fixed_response[:100]}...")
            return fixed_response
    except Exception as e:
        error_msg = str(e)
        print(f"Error in Groq API call: {error_msg}")
        error_detail = "An error occurred during Groq processing."
        if "401" in error_msg or "invalid_api_key" in error_msg.lower():
            error_detail = "Invalid Groq API key. Check your .env file and get a valid key from https://console.groq.com"
        elif "429" in error_msg or "rate limit" in error_msg.lower():
            error_detail = "Groq rate limit exceeded. Please try again later."
        return f'{{"name": "Groq Processing Error", "email": "Not provided", "phone_number": "Not provided", "experience": "{error_detail}", "skills": [], "cultural_fit": "Unknown", "communication": "Unknown", "ranking_score": 0, "explanation": "{error_detail}"}}'

@app.route('/process', methods=['POST'])
@auth_required
def process_resumes():
    if 'resumes' not in request.files:
        return jsonify({"error": "No files uploaded"}), 400
    job_title = request.form.get("job_title", "")
    job_description = request.form.get("job_description", "")
    files = request.files.getlist("resumes")
    if not job_title:
        return jsonify({"error": "Job title is required"}), 400
    print(f"Processing {len(files)} resumes for job: {job_title}")
    print(f"Job description length: {len(job_description)} characters")
    results = []
    resume_paths = {}
    processed_count = 0
    error_count = 0
    if len(files) == 0:
        print("ERROR: No files received in request")
        return jsonify({"error": "No files uploaded"}), 400
    ALLOWED_EXTENSIONS = {'.pdf', '.doc', '.docx', '.jpg', '.jpeg', '.png'}
    for i, file in enumerate(files):
        print(f"Processing file {i}: filename='{file.filename}', content_type='{file.content_type}'")
        if file.filename == '':
            print(f"WARNING: File {i} has empty filename, skipping")
            continue
        filename_lower = file.filename.lower()
        file_extension = os.path.splitext(filename_lower)[1]
        if file_extension not in ALLOWED_EXTENSIONS:
            print(f"WARNING: File {i} has invalid extension '{file_extension}', skipping")
            error_count += 1
            results.append({
                "name": "Invalid File Type",
                "email": "Not provided",
                "phone_number": "Not provided",
                "experience": "Unsupported file format",
                "skills": [],
                "cultural_fit": "Unknown",
                "communication": "Unknown",
                "ranking_score": 0,
                "explanation": f"File type '{file_extension}' is not supported. Please upload PDF, DOC, DOCX, JPG, or PNG files only.",
                "file": file.filename,
                "file_index": i
            })
            continue
        filename = file.filename
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(filepath)
        resume_paths[filename] = filepath
        resume_paths[str(i)] = filepath
        resume_paths[i] = filepath
        print(f"Saved file {i}: {filename} to {filepath}")
        resume_text = extract_resume_text(filepath)
        if not resume_text:
            print(f"ERROR: Could not extract text from resume: {filename}")
            print(f"File exists: {os.path.exists(filepath)}, File size: {os.path.getsize(filepath) if os.path.exists(filepath) else 'N/A'} bytes")
            error_count += 1
            continue
        print(f"Extracted {len(resume_text)} characters from resume: {filename}")
        if i > 0:
            delay = 3.0 + (random.random() * 2.0)
            print(f"Adding delay of {delay:.2f}s before processing next resume to avoid rate limits...")
            time.sleep(delay)
        prompt = f"""You are an AI HR evaluator. Your task is to review a candidate's resume for a specific job position and provide a detailed assessment in JSON format.

Given:
- Job Title: {job_title}
- Job Requirements: {job_description} (if provided)
- Candidate Resume: {resume_text}

Instructions:
1. *Extract Candidate's Name:*
   - Find and record the candidate's name from the resume. It's usually at the top or in the header.

2. *Extract Contact Information:*
   - Find and record the candidate's email address and phone number from the resume. These are typically near the name or in a contact section. Use "Not provided" if not found.

3. *Analyze Experience:*
   - From the resume, identify the work experience section and calculate the total years of relevant experience.
   - Identify roles, positions, or projects that are directly relevant to the job title or job requirements.
   - Focus on skills demonstrated, responsibilities handled, and achievements accomplished rather than employment gaps.
   - Write the experience summary in natural, human-friendly language — as if a real recruiter wrote it. Avoid bullet-point labels like 'Summary:', 'Total years of experience:', or 'Relevant experience projects:'. Instead, weave the details into one or two flowing, natural paragraphs that sound like a person describing the candidate's background.

4. *Identify Skills and Competencies:*
   - From the resume, identify both explicit skills (listed in skills section) and implicit skills (demonstrated through accomplishments, projects, or responsibilities).
   - Consider technical skills, soft skills, and transferable skills that could be valuable for the position.
   - Prioritize skills that align with the job requirements or are typical for the job title if requirements aren't specified.

5. *Evaluate Fit:*
   - *Step 1: Identify Core Job Requirements:*
     - If job requirements are provided, extract key qualifications, skills, and experience needs.
     - If only job title is provided, infer standard qualifications and skills typically expected for that role (e.g., Python, machine learning for Data Scientist).
     - Distinguish between truly essential requirements (must-haves) and "nice-to-have" qualifications.
     - Avoid treating degree requirements as essential unless they represent regulatory or legal requirements (e.g., medical license for doctors).

    - *Step 2: Assess Skill Alignment (50% of score):*
      - Count the exact number of essential skills required for the role versus the number the candidate demonstrably possesses.
      - award points ONLY for skills clearly evidenced in the resume (stated or demonstrated through projects/achievements). Do NOT assume skills not present.
      - Calculate: skill_match_ratio = demonstrated_essential_skills / total_essential_skills.
      - skill_alignment_score = skill_match_ratio * 50, rounded to the nearest integer.
      - Apply strict penalties:
        * If skill_match_ratio >= 0.9: score 40-50
        * If skill_match_ratio >= 0.7: score 30-39
        * If skill_match_ratio >= 0.5: score 20-29
        * If skill_match_ratio >= 0.3: score 10-19
        * If skill_match_ratio < 0.3: score 0-9
      - Missing 3 or more core/essential skills => maximum 10 points for this section.
      - If fewer than 40% of essential skills are met, cap the total_ranking_score at 25.

    - *Step 3: Evaluate Experience Quality (40% of score):*
      - Assess ONLY directly relevant experience to the job title and requirements. Irrelevant experience earns 0 points here.
      - Judge depth: years of relevant experience, complexity of responsibilities, measurable achievements, and evidence of progression.
      - Base points on years AND impact:
        * 5+ years of directly relevant experience with major achievements: 32-40
        * 3-5 years of relevant experience with solid achievements: 24-31
        * 1-3 years of relevant experience or some relevant projects: 16-23
        * Less than 1 year of relevant experience or internships only: 8-15
        * No relevant experience at all: 0-7
      - If the candidate's experience is entirely unrelated to the role, assign 0 points.

    - *Step 4: Assess Growth Potential (5% of score):*
      - Award 0-5 points based on evidence of continuous learning (certifications, new skills, career progression).
      - Only award 4-5 if there is clear, recent growth. Award 0 if no evidence.

    - *Step 5: Evaluate Unique Value (5% of score):*
      - Award 0-5 points for exceptional combinations: rare skills, leadership, startup experience, open-source contributions, awards.
      - Most candidates should score 0-2 here. Only truly standout candidates get 4-5.

    - *Step 6: Calculate Overall Ranking Score:*
      - total_ranking_score = skill_alignment_score + experience_quality_score + growth_potential_score + unique_value_score

    - *CRITICAL SCORING RULES - BREAK CLUSTERING:*
      - YOU MUST USE THE FULL 0-100 RANGE. Do NOT cluster most candidates between 75-90.
      - STRICT calibration examples:
        * 95-100: Rare - perfect skill match (100% essential skills), 7+ years of impactful experience, strong evidence of leadership/innovation. Top 1% of candidates.
        * 85-94: Excellent - near-perfect skill match (>=90% essential skills), 5+ years strong experience, clear achievements. Top 5%.
        * 70-84: Good - solid skill match (70-89% essential skills), 3+ years relevant experience, some achievements. Above average.
        * 55-69: Average - partial skill match (50-69% essential skills), 1-3 years experience, limited achievements. Typical candidate.
        * 35-54: Below average - poor skill match (30-49% essential skills), minimal experience or frequent job changes. Significant gaps.
        * 15-34: Weak - very poor skill match (15-29% essential skills), almost no relevant experience. Major red flags.
        * 0-14: Unacceptable - fails to meet minimum requirements (<15% essential skills), no relevant experience. Do not hire.
      - Force differentiation: If two candidates differ in skill match by 20% or more, their final scores MUST differ by at least 10 points.
      - Penalize incomplete profiles heavily: if critical information is missing (e.g., no contact info, no experience details), reduce score by 10-15 points.
      - Cap absolute maximum at 100 and absolute minimum at 0.

6. *Provide Explanation:*
   - Give a balanced explanation of the score, highlighting strengths and growth areas.
   - Emphasize skill and experience alignment as primary factors, with growth potential and unique value as secondary.
   - Acknowledge alternative paths to skills acquisition beyond traditional routes.

7. *Bias Prevention Guidelines:*
   - Do not consider or mention factors such as name origins, gender, ethnicity, race, religion, age, or personal details.
   - Evaluate gaps in employment history neutrally, without negative assumptions.
   - Do not favor prestigious institutions or discount learning from non-traditional education paths.
   - Focus on demonstrated capabilities rather than proxies like specific degrees.
   - Use neutral language and avoid terms with gendered or cultural connotations.
   - Evaluate international experience and qualifications fairly, recognizing equivalent skills across different systems.

Output Format:
Return only a valid JSON object with the following keys, no additional text:
- "name": "The candidate's full name, if detectable."
- "email": "The candidate's email address, or 'Not provided' if not found."
- "phone_number": "The candidate's phone number, or 'Not provided' if not found."
- "experience": "A summary in precise bullet points of the candidate's years of experience, key roles, and notable achievements."
- "skills": "A list of top skills demonstrated in the resume relevant to the job."
- "cultural_fit": "An evaluation of the candidate's potential cultural contribution based on their professional values and work approach."
- "communication": "An assessment of the candidate's communication skills as evidenced by their resume (e.g., clarity, professionalism)."
- "ranking_score": "A numerical fit score (0-100) that reflects the overall suitability of the candidate for the role relative to other resumes."
- "explanation": "A concise explanation of why this candidate is a good or less ideal fit for the role, highlighting strengths and any potential growth areas."

Return ONLY a valid JSON object with these exact keys:
{{
  "name": "string",
  "email": "string",
  "phone_number": "string",
  "experience": "string",
  "skills": ["array", "of", "strings"],
  "cultural_fit": "string",
  "communication": "string",
  "ranking_score": number,
  "explanation": "string"
}}

Rules:
1. No markdown formatting
2. No additional text outside the JSON
3. No thinking process or commentary
4. All strings must be properly quoted
5. Ensure all brackets are properly closed"""
        try:
            response = None
            if groq_client is not None:
                print(f"Calling Groq API for file {filename}...")
                response = call_groq(prompt)
                if response:
                    print(f"✓ Groq processed {filename} successfully")
            else:
                print(f"ERROR: Groq client is unavailable")
                error_count += 1
                results.append({
                    "name": "Processing Failed",
                    "email": "Not provided",
                    "phone_number": "Not provided",
                    "experience": "No AI service available",
                    "skills": [],
                    "cultural_fit": "Unknown",
                    "communication": "Unknown",
                    "ranking_score": 0,
                    "explanation": "Groq API key is not configured. Check your .env file.",
                    "file": filename,
                    "file_index": i
                })
                continue
            if not response:
                print(f"ERROR: No response from AI for file {filename}")
                error_count += 1
                results.append({
                    "name": "Processing Failed",
                    "email": "Not provided",
                    "phone_number": "Not provided",
                    "experience": "Failed to process resume",
                    "skills": [],
                    "cultural_fit": "Unknown",
                    "communication": "Unknown",
                    "ranking_score": 0,
                    "explanation": "The AI service failed to process this resume. Please check your API keys.",
                    "file": filename,
                    "file_index": i
                })
                continue
            response = response.strip()
            if response.startswith('```json'):
                response = response[7:]
            if response.endswith('```'):
                response = response[:-3]
            response = response.strip()
            try:
                candidate_info = json.loads(response)
                if "candidateName" in candidate_info and "name" not in candidate_info:
                    candidate_info["name"] = candidate_info["candidateName"]
                if "candidate_name" in candidate_info and "name" not in candidate_info:
                    candidate_info["name"] = candidate_info["candidate_name"]
                if "emailAddress" in candidate_info and "email" not in candidate_info:
                    candidate_info["email"] = candidate_info["emailAddress"]
                if "candidate_email" in candidate_info and "email" not in candidate_info:
                    candidate_info["email"] = candidate_info["candidate_email"]
                if "phoneNumber" in candidate_info and "phone_number" not in candidate_info:
                    candidate_info["phone_number"] = candidate_info["phoneNumber"]
                if "phone" in candidate_info and "phone_number" not in candidate_info:
                    candidate_info["phone_number"] = candidate_info["phone"]
                if "score" in candidate_info and "ranking_score" not in candidate_info:
                    candidate_info["ranking_score"] = candidate_info["score"]
                if "overall_score" in candidate_info and "ranking_score" not in candidate_info:
                    candidate_info["ranking_score"] = candidate_info["overall_score"]
                if "skills" in candidate_info and not isinstance(candidate_info["skills"], list):
                    if isinstance(candidate_info["skills"], str):
                        candidate_info["skills"] = [skill.strip() for skill in candidate_info["skills"].split(",")]
                    else:
                        candidate_info["skills"] = []
                candidate_info["file"] = filename
                candidate_info["file_index"] = i
                candidate_name = candidate_info.get('name', 'Unknown')
                print(f"✓ Successfully processed candidate from file {filename}: {candidate_name}")
                print(f"  - Email: {candidate_info.get('email', 'Not provided')}")
                print(f"  - Ranking Score: {candidate_info.get('ranking_score', 0)}")
                results.append(candidate_info)
                processed_count += 1
            except json.JSONDecodeError as e:
                print(f"JSON error processing file {filename}: {str(e)}")
                match = re.search(r'\{.*\}', response, re.DOTALL)
                if match:
                    candidate_info = json.loads(match.group(0))
                    candidate_info["file"] = filename
                    results.append(candidate_info)
                error_count += 1
            except Exception as e:
                print(f"Error processing candidate info for file {filename}: {str(e)}")
                traceback.print_exc()
                error_count += 1
        except Exception as e:
            print(f"Error during AI processing of file {filename}: {str(e)}")
            traceback.print_exc()
            error_count += 1
    if not results:
        return jsonify({"error": "No valid candidate data processed", "details": "All resume processing attempts failed"}), 400
    try:
        print(f"Processing batch of {len(results)} candidates...")
        from utils.supabase_db import test_supabase_connection
        db_ok, db_msg = test_supabase_connection()
        if not db_ok:
            print(f"WARNING: Supabase unavailable ({db_msg}). Using local storage fallback.")
            print(f"Processing {len(results)} candidates with local storage...")
        else:
            print(f"✓ Supabase connection OK. Using cloud storage.")
        user_id = request.session.get('sub', '') if hasattr(request, 'session') else ''
        processed_results = CandidateProcessor.process_candidates_batch(results, job_role=job_title, resume_paths=resume_paths, user_id=user_id)
        print(f"Batch processing complete. Got {len(processed_results)} processed candidates.")
        if processed_results:
            simplified_results = []
            for candidate in processed_results:
                simplified_results.append({
                    "id": candidate.get("id"),
                    "name": candidate.get("name"),
                    "email": candidate.get("email"),
                    "phone_number": candidate.get("phone_number"),
                    "ranking_score": candidate.get("ranking_score"),
                    "resume_link": candidate.get("resume_link"),
                    "experience": candidate.get("experience"),
                    "skills": candidate.get("skills"),
                    "cultural_fit": candidate.get("cultural_fit"),
                    "communication": candidate.get("communication"),
                    "explanation": candidate.get("explanation"),
                    "file": candidate.get("file") or candidate.get("resume_link")
                })
            simplified_results.sort(key=lambda x: float(x.get('ranking_score', 0)), reverse=True)
            return jsonify({"status": "success", "results": simplified_results, "count": len(simplified_results), "processed_count": processed_count, "error_count": error_count, "job_title": job_title})
        else:
            error_msg = "Failed to process candidates - database operation failed"
            debug_info = {"supabase_initialized": supabase is not None, "results_count": len(results), "processed_results_count": len(processed_results) if 'processed_results' in locals() else 0}
            if supabase is None:
                error_msg += ". Supabase client is not initialized. Check your .env file for SUPABASE_URL and SUPABASE_KEY."
            else:
                from utils.supabase_db import test_supabase_connection
                db_ok, db_msg = test_supabase_connection()
                debug_info["db_connection_test"] = {"success": db_ok, "message": db_msg}
            return jsonify({"error": error_msg, "processed_count": processed_count, "error_count": error_count, "details": "Check server logs for more information. The resume was processed by AI but failed to save to database.", "debug": debug_info}), 500
    except Exception as e:
        print(f"Error in final processing: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": f"Failed to process results: {str(e)}", "processed_count": processed_count, "error_count": error_count}), 500

@app.route('/send-interview-invitation', methods=['POST'])
def send_interview_invitation():
    """
    Send interview invitation email via backend using Brevo SMTP API
    Email body is generated using Gemini AI if available, falls back to static template.
    """
    try:
        data = request.get_json(silent=True) or {}
        candidate_email = data.get('candidateEmail')
        candidate_name = data.get('candidateName')
        venue = data.get('venue')
        time_slot = data.get('timeSlot')

        if not candidate_email or not candidate_name or not venue or not time_slot:
            return jsonify({"error": "Missing required email fields"}), 400

        brevo_api_key = os.getenv('BREVO_API_KEY', '')
        if brevo_api_key.startswith('"') and brevo_api_key.endswith('"'):
            brevo_api_key = brevo_api_key[1:-1]
        if brevo_api_key.startswith("'") and brevo_api_key.endswith("'"):
            brevo_api_key = brevo_api_key[1:-1]
        print(f"BREVO_API_KEY loaded: prefix={brevo_api_key[:12]}... len={len(brevo_api_key)}")
        if not brevo_api_key:
            print("BREVO_API_KEY missing after cleanup")
            return jsonify({"error": "Email service is not configured on the server."}), 500

        # Generate email body using Gemini
        email_body = None
        if genai_client and GEMINI_API_KEY:
            try:
                gemini_model = genai.GenerativeModel("gemini-2.0-flash-lite")
                gemini_prompt = f"""You are an HR professional. Generate a formal, polished interview invitation email in HTML format using clean inline styles. The tone must be professional, courteous, and warm but not casual.

Use this information:
- Candidate Name: {candidate_name}
- Venue: {venue}
- Time Slot: {time_slot}

Structure the email as follows:
1. Subject line inside an <h1> tag: "Interview Invitation - TechCorp"
2. Formal salutation: "Dear Mr./Ms. {candidate_name},"
3. A short paragraph expressing that the candidate's qualifications stood out and TechCorp would like to invite them for an interview at the specified venue and time.
4. A clear section with the venue and time details in a clean layout.
5. A polite request to confirm availability by replying to this email.
6. A professional closing: "Best regards," followed by "HR Team" and "TechCorp" on separate lines.

Return ONLY the raw HTML code. No markdown, no code fences, no backticks, no CSS style tags — only inline styles on each HTML element. Do not wrap the output in any additional text or formatting."""
                gemini_response = gemini_model.generate_content(gemini_prompt)
                if gemini_response and gemini_response.text:
                    email_body = gemini_response.text.strip()
                    # Clean up any accidental markdown fences from the response
                    if email_body.startswith('```html'):
                        email_body = email_body[7:]
                    elif email_body.startswith('```'):
                        email_body = email_body[3:]
                    if email_body.endswith('```'):
                        email_body = email_body[:-3]
                    email_body = email_body.strip()
                    print(f"✓ Gemini generated email body ({len(email_body)} chars)")
            except Exception as gemini_err:
                print(f"Gemini email generation failed: {gemini_err}")
                email_body = None

        # Fallback to static professional template if Gemini failed
        if not email_body:
            company_name = "TechCorp"
            email_body = f"""
            <html>
              <body style="font-family: 'Segoe UI', Arial, sans-serif; background-color: #f4f4f4; margin: 0; padding: 20px;">
                <table width="100%" cellpadding="0" cellspacing="0" style="max-width: 600px; margin: 0 auto; background-color: #ffffff; border-radius: 8px; overflow: hidden; box-shadow: 0 2px 8px rgba(0,0,0,0.1);">
                  <tr>
                    <td style="background: linear-gradient(135deg, #1a5276, #2980b9); padding: 30px 20px; text-align: center;">
                      <h1 style="color: #ffffff; font-size: 26px; margin: 0; letter-spacing: 1px;">INTERVIEW INVITATION</h1>
                    </td>
                  </tr>
                  <tr>
                    <td style="padding: 30px 30px 20px 30px;">
                      <p style="font-size: 16px; color: #333333; margin: 0 0 5px 0;">Dear <strong style="color: #1a5276;">{candidate_name}</strong>,</p>
                      <p style="font-size: 15px; color: #555555; line-height: 1.6; margin: 15px 0 0 0;">
                        Thank you for your interest in joining <strong>{company_name}</strong>. After reviewing your application and qualifications, we are pleased to invite you for an interview at our facility. Your background and experience align well with what we are looking for, and we look forward to discussing your potential contribution to our team.
                      </p>
                    </td>
                  </tr>
                  <tr>
                    <td style="padding: 10px 30px;">
                      <table width="100%" cellpadding="0" cellspacing="0" style="background-color: #eaf2f8; border-radius: 6px; border-left: 4px solid #2980b9;">
                        <tr>
                          <td style="padding: 20px;">
                            <p style="font-size: 16px; color: #1a5276; font-weight: bold; margin: 0 0 10px 0;">Interview Details</p>
                            <p style="font-size: 14px; color: #333333; margin: 5px 0;"><strong>Venue:</strong> {venue}</p>
                            <p style="font-size: 14px; color: #333333; margin: 5px 0;"><strong>Time:</strong> {time_slot}</p>
                            <p style="font-size: 14px; color: #333333; margin: 5px 0 0 0;"><strong>Duration:</strong> Approximately 45-60 minutes</p>
                          </td>
                        </tr>
                      </table>
                    </td>
                  </tr>
                  <tr>
                    <td style="padding: 20px 30px 30px 30px;">
                      <p style="font-size: 15px; color: #555555; line-height: 1.6; margin: 0;">
                        Kindly confirm your availability for the scheduled time by replying to this email at your earliest convenience. If you have any questions or require further information, please do not hesitate to reach out.
                      </p>
                      <p style="font-size: 15px; color: #555555; line-height: 1.6; margin: 15px 0 0 0;">
                        We look forward to meeting you.
                      </p>
                    </td>
                  </tr>
                  <tr>
                    <td style="border-top: 1px solid #e0e0e0; padding: 20px 30px;">
                      <p style="font-size: 15px; color: #333333; margin: 0 0 3px 0;">Best regards,</p>
                      <p style="font-size: 16px; color: #1a5276; font-weight: bold; margin: 0 0 3px 0;">HR Team</p>
                      <p style="font-size: 14px; color: #777777; margin: 0;">{company_name}</p>
                    </td>
                  </tr>
                </table>
              </body>
            </html>
            """
            print("Using static professional template (Gemini unavailable or failed)")

        payload = {
            "sender": {"name": "HR Team", "email": "ushagpt1372@gmail.com"},
            "to": [{"email": candidate_email, "name": candidate_name}],
            "subject": "Final Interview Invitation - TechCorp",
            "htmlContent": email_body
        }

        import requests as http_requests
        response = http_requests.post(
            "https://api.brevo.com/v3/smtp/email",
            headers={
                "accept": "application/json",
                "content-type": "application/json",
                "api-key": brevo_api_key
            },
            json=payload,
            timeout=15
        )

        try:
            resp_json = response.json()
        except Exception:
            resp_json = {}

        if response.ok and response.status_code in (200, 201):
            return jsonify({"status": "success", "message": "Email sent successfully"}), 200
        else:
            error_message = resp_json.get('message') or resp_json.get('error') or resp_json or 'Failed to send email'
            print(f"Brevo email send failed: {response.status_code} {error_message}")
            status_label = 'authentication/config error' if response.status_code in (401, 403) else 'provider error'
            return jsonify({"error": f"Email service {status_label}: {error_message}"}), 500

    except Exception as e:
        print(f"Error sending interview invitation: {str(e)}")
        traceback.print_exc()
        return jsonify({"error": "Failed to send email"}), 500

@app.route('/test-db', methods=['GET'])
def test_database():
    from utils.supabase_db import test_supabase_connection, store_job_role, store_candidate
    from utils.supabase_client import initialize_supabase
    results = {"supabase_initialized": supabase is not None, "tests": {}}
    if supabase:
        db_ok, db_msg = test_supabase_connection()
        results["tests"]["connection"] = {"success": db_ok, "message": db_msg}
    else:
        results["tests"]["connection"] = {"success": False, "message": "Supabase client not initialized"}
    if supabase:
        try:
            test_query = supabase.table('open_roles').select('id').limit(1).execute()
            results["tests"]["query_roles"] = {"success": True, "message": f"Found {len(test_query.data) if test_query.data else 0} roles"}
        except Exception as e:
            results["tests"]["query_roles"] = {"success": False, "message": str(e)}
    if supabase:
        try:
            test_query = supabase.table('candidates').select('id').limit(1).execute()
            results["tests"]["query_candidates"] = {"success": True, "message": f"Found {len(test_query.data) if test_query.data else 0} candidates"}
        except Exception as e:
            results["tests"]["query_candidates"] = {"success": False, "message": str(e)}
    if supabase:
        try:
            test_role_name = f"TEST_ROLE_{int(time.time())}"
            role_result = store_job_role(test_role_name)
            if role_result:
                results["tests"]["insert_role"] = {"success": True, "message": f"Created test role: {role_result}"}
            else:
                results["tests"]["insert_role"] = {"success": False, "message": "store_job_role returned None"}
        except Exception as e:
            results["tests"]["insert_role"] = {"success": False, "message": str(e)}
    if supabase:
        try:
            test_candidate = {"name": "Test Candidate", "email": f"test_{int(time.time())}@test.com", "phone_number": "123-456-7890"}
            candidate_id = store_candidate(test_candidate)
            if candidate_id:
                results["tests"]["insert_candidate"] = {"success": True, "message": f"Created test candidate with ID: {candidate_id}"}
            else:
                results["tests"]["insert_candidate"] = {"success": False, "message": "store_candidate returned None"}
        except Exception as e:
            results["tests"]["insert_candidate"] = {"success": False, "message": str(e)}
    return jsonify(results)

@app.route('/roles', methods=['GET'])
@auth_required
def get_roles():
    try:
        user_id = request.session.get('sub', '') if hasattr(request, 'session') else ''
        print(f"Fetching roles for user: {user_id}")
        roles = get_all_roles(user_id)
        return jsonify({"status": "success", "results": roles, "count": len(roles)})
    except Exception as e:
        return jsonify({"error": f"Failed to get roles: {str(e)}"}), 500

@app.route('/candidates_for_role/<role_name>', methods=['GET'])
@auth_required
def get_candidates_by_role(role_name):
    try:
        user_id = request.session.get('sub', '') if hasattr(request, 'session') else ''
        print(f"Fetching candidates for role: {role_name} for user: {user_id}")
        candidates = get_candidates_for_role(role_name, user_id)
        return jsonify({"status": "success", "results": candidates, "count": len(candidates), "role_name": role_name})
    except Exception as e:
        return jsonify({"error": f"Failed to get candidates for role {role_name}: {str(e)}"}), 500

@app.route('/roles/<role_name>', methods=['DELETE'])
@auth_required
def delete_role_route(role_name):
    try:
        user_id = request.session.get('sub', '') if hasattr(request, 'session') else ''
        success = False
        if supabase:
            try:
                role_query = supabase.table('open_roles').select('id').eq('role_name', role_name)
                if user_id:
                    role_query = role_query.eq('user_id', user_id)
                role_resp = role_query.execute()
                if role_resp.data:
                    role_id = role_resp.data[0]['id']
                    candidates_query = supabase.table('role_candidates').select('candidate_id').eq('role_id', role_id)
                    if user_id:
                        candidates_query = candidates_query.eq('user_id', user_id)
                    candidates_resp = candidates_query.execute()
                    if candidates_resp.data:
                        for rc in candidates_resp.data:
                            candidate_id = rc['candidate_id']
                            c_resp = supabase.table('candidates').select('resume_link,user_id').eq('id', candidate_id).execute()
                            if c_resp.data and len(c_resp.data) > 0:
                                cd = c_resp.data[0]
                                resume_link = cd.get('resume_link', '')
                                cand_user_id = cd.get('user_id', '')
                                if resume_link and not resume_link.startswith('file://'):
                                    try:
                                        file_name = resume_link.rsplit('/', 1)[-1]
                                        if file_name:
                                            from utils.supabase_storage import delete_resume
                                            delete_resume(file_name)
                                            print(f"Deleted resume '{file_name}' from storage for candidate {candidate_id}")
                                    except Exception as storage_error:
                                        print(f"Warning: Could not delete resume from storage: {storage_error}")
                                try:
                                    c_del = supabase.table('candidates').delete().eq('id', candidate_id)
                                    if user_id and cand_user_id:
                                        c_del = c_del.eq('user_id', user_id)
                                    c_del.execute()
                                    print(f"Deleted candidate {candidate_id} from candidates table")
                                except Exception as cand_del_err:
                                    print(f"Warning: Could not delete candidate {candidate_id}: {cand_del_err}")
                    rc_query = supabase.table('role_candidates').delete().eq('role_id', role_id)
                    if user_id:
                        rc_query = rc_query.eq('user_id', user_id)
                    rc_query.execute()
                    supabase.table('open_roles').delete().eq('id', role_id).execute()
                    success = True
            except Exception as del_err:
                print(f"Supabase delete failed: {del_err}")
        if not success:
            success = delete_role_local(role_name, user_id)
        if success:
            return jsonify({"status": "success", "message": f"Role '{role_name}' deleted"})
        else:
            return jsonify({"error": f"Failed to delete role '{role_name}'"}), 500
    except Exception as e:
        print(f"Error deleting role '{role_name}': {str(e)}")
        return jsonify({"error": str(e)}), 500

@app.route('/candidates/<int:candidate_id>', methods=['DELETE'])
@auth_required
def delete_candidate_route(candidate_id):
    try:
        user_id = request.session.get('sub', '') if hasattr(request, 'session') else ''
        success = delete_candidate(candidate_id, user_id)
        if not success:
            success = delete_candidate_local(candidate_id, user_id)
        if success:
            return jsonify({"status": "success", "message": f"Candidate {candidate_id} deleted"})
        else:
            return jsonify({"error": f"Failed to delete candidate {candidate_id}"}), 500
    except Exception as e:
        print(f"Error deleting candidate {candidate_id}: {str(e)}")
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)